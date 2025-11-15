import os
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
from scipy.stats import t
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

class RentalAnalyzer():
    def __init__(self, params):
        # Initialize the RentalAnalyzer class with the provided parameters
        self.df_list = params.get('df_list')
        self.city_list = params.get('city_list')

        # Paths for generated files
        self.images_path = 'images'
        self.reports_path = 'reports'
        
        # Attributes to store analysis results
        self.stats_dict = {}
        self.analysis_data = {} # Will store the text report for each city
        self.city_ranking = []  # Will store the final ranked list

        # Ensure directories exist
        if not os.path.exists(self.images_path):
            os.makedirs(self.images_path)
        if not os.path.exists(self.reports_path):
            os.makedirs(self.reports_path)
    
    # Calculate statistical measures for a given DataFrame column
    def _calculate_stats(self, df, column='price_per_sqm', conf_level=0.95):
        sample_size = len(df)
        mean = df[column].mean()
        median = df[column].median()
        mode = df[column].mode()
        skewness = df[column].skew()
        variance = df[column].var()
        std_dev = df[column].std()
        coefficient_of_variation = std_dev / mean
        dfg = len(df) - 1
        crit_value = t.ppf((1 + conf_level) / 2, dfg)
        lower_bound = mean - crit_value * (std_dev / np.sqrt(len(df)))
        upper_bound = mean + crit_value * (std_dev / np.sqrt(len(df)))
        stats = {
            'sample_size': sample_size,
            'mean': mean,
            'median': median,
            'mode': mode,
            'skewness': skewness,
            'variance': variance,
            'standard_deviation': std_dev,
            'coefficient_of_variation': coefficient_of_variation,
            'conf_level': conf_level,
            'lower_bound': lower_bound,
            'upper_bound': upper_bound
        }
        return stats

    # Append calculated stats to the stats dictionary
    def _append_stats(self, key, stats):
        for column, value in stats.items():
            self.stats_dict[key][column].append(value)
        
    # Populate the stats dictionary for all DataFrames
    def _get_stats_dict(self):
        self.stats_dict = {
            'rent': {'sample_size': [], 'mean': [], 'median': [], 'mode': [], 'skewness': [], 'variance': [],
                     'standard_deviation': [], 'coefficient_of_variation': [], 'conf_level': [], 'lower_bound': [], 
                     'upper_bound': []},
            'sell': {'sample_size': [], 'mean': [], 'median': [], 'mode': [], 'skewness': [], 'variance': [],
                     'standard_deviation': [], 'coefficient_of_variation': [], 'conf_level': [], 'lower_bound': [], 
                     'upper_bound': []}
        }
        for df in self.df_list:
            key = 'rent' if 'rent' in df.name else 'sell'
            stats = self._calculate_stats(df)
            self._append_stats(key, stats)
    
    # Generate and save price per square meter plots for a city
    def _plot_price_per_sqm(self, city, sell_df, rent_df):
        n_cols = 2
        n_rows = 1
        fig, axs = plt.subplots(n_rows, n_cols, figsize=(15, 4))

        sell_ax = axs[0]
        rent_ax = axs[1]
        sell_ax.set_title(f'{city} (Sell)')
        rent_ax.set_title(f'{city} (Rent)')

        sell_counts, sell_bins, _ = sell_ax.hist(sell_df['price_per_sqm'], bins=15, edgecolor='black', linewidth=1)
        sell_ticks = np.linspace(sell_bins[0], sell_bins[-1], num=15, endpoint=True)
        sell_ax.set_xticks(sell_ticks)

        rent_counts, rent_bins, _ = rent_ax.hist(rent_df['price_per_sqm'], bins=15, edgecolor='black', linewidth=1)
        rent_ticks = np.linspace(rent_bins[0], rent_bins[-1], num=15, endpoint=True)
        rent_ax.set_xticks(rent_ticks)

        plt.tight_layout()
        
        plot_filename = os.path.join(self.images_path, f'{city}_price_per_sqm.png')
        plt.savefig(plot_filename)
        print(f"Saved plot to {plot_filename}")
        plt.close(fig) 
    
    # Estimate monthly return based on selling and renting prices
    def _estimate_monthly_return(self, selling_price_per_sqm, renting_price_per_sqm, property_value=100000):
        property_size = property_value / selling_price_per_sqm
        monthly_rental_income = renting_price_per_sqm * property_size
        monthly_return = (monthly_rental_income / property_value) * 100
        return monthly_return

    # HELPER METHOD
    # This method holds the text block
    def _build_city_report_string(self, city, i, lower_estimate, upper_estimate):
        """Builds the formatted text block for a city's analysis."""
        sell = self.stats_dict['sell']
        rent = self.stats_dict['rent']
        
        text = f"""
Based on the data from storia.ro, the following observations can be made for {city}:

1. Analysis of property listed for sale:

Sample Size: {sell['sample_size'][i]} estates listed for sale.
Average Selling Price per Square Meter: {sell['mean'][i]:.2f}
Median Selling Price per Square Meter: {sell['median'][i]:.2f}
Most Common Selling Price per Square Meter: {sell['mode'][i][0]}
Skewness: {sell['skewness'][i]:.2f}
Standard deviation: {sell['standard_deviation'][i]:.2f}
Coefficient of Variation: {sell['coefficient_of_variation'][i]:.2f}
Confidence intervals: {sell['conf_level'][i]*100}% confidence that the true average price of a 
square meter is between {sell['lower_bound'][i]:.2f} - {sell['upper_bound'][i]:.2f}


2. Analysis of property listed for rent:

Sample Size: {rent['sample_size'][i]} estates listed for rent.
Average Renting Price per Square Meter: {rent['mean'][i]:.2f}
Median Renting Price per Square Meter: {rent['median'][i]:.2f}
Most Common Renting Price per Square Meter: {rent['mode'][i][0]}
Skewness: {rent['skewness'][i]:.2f}
Standard deviation: {rent['standard_deviation'][i]:.2f}
Coefficient of Variation: {rent['coefficient_of_variation'][i]:.2f}
Confidence intervals: {rent['conf_level'][i]*100}% confidence that the true average price of a 
square meter is between {rent['lower_bound'][i]:.2f} - {rent['upper_bound'][i]:.2f}

For every euro spent buying a square meter in {city}, we can expect a monthly income from rent between
{lower_estimate:.2f} and {upper_estimate:.2f} euro.
"""
        return text

    # This method performs calculations and saves results to the class
    def analyzer(self):
        """
        Runs all calculations, generates plots, and stores results
        in self.analysis_data and self.city_ranking.
        """
        self._get_stats_dict()
        sell_dfs = [df for df in self.df_list if 'sell' in df.name]
        rent_dfs = [df for df in self.df_list if 'rent' in df.name]
        sell = self.stats_dict['sell']
        rent = self.stats_dict['rent']
        
        city_ranking_data = [] # Temporary list
        
        for city, sell_df, rent_df, i in zip(self.city_list, sell_dfs, rent_dfs, list(range(len(self.city_list)+1))): 
            lower_estimate = self._estimate_monthly_return(sell['lower_bound'][i], rent['lower_bound'][i])
            upper_estimate = self._estimate_monthly_return(sell['upper_bound'][i], rent['upper_bound'][i])
            
            # Generate and save the plot
            self._plot_price_per_sqm(city, sell_df, rent_df)
            
            # Generate text block and save to class attribute
            self.analysis_data[city] = self._build_city_report_string(city, i, lower_estimate, upper_estimate)
            
            # Append data for final ranking
            city_ranking_data.append((city, lower_estimate, upper_estimate))
        
        # Sort and save final ranking to class attribute
        self.city_ranking = sorted(city_ranking_data, key=lambda x: x[1], reverse=True)
        print("Analysis and plot generation complete.") # Log to console

    # METHOD for console output
    def print_console_report(self):
        """
        Prints the pre-calculated analysis results to the console.
        """
        if not self.analysis_data or not self.city_ranking:
            print("Error: Analysis has not been run. Call analyzer() first.")
            return

        for city in self.city_list:
            print(self.analysis_data[city])
        
        print("\nRanking based on monthly rental income estimates:\n")
        for rank, (city, lower, upper) in enumerate(self.city_ranking, start=1):
            print(f"""{rank}. {city}: 
The expected monthly income from rent for every euro spent is estimated to be between {lower:.2f} and {upper:.2f}.\n""")

    # This method presents data to the DOCX file
    def generate_word_report(self, report_filename="Romanian_Apartment_Analysis_Report.docx"):
        """
        Saves the pre-calculated results and plots to a Word document.
        """
        if not self.analysis_data or not self.city_ranking:
            print("Error: Analysis has not been run. Call analyzer() first.")
            return
            
        full_report_path = os.path.join(self.reports_path, report_filename)
        print(f"\nGenerating Word report: {full_report_path}...")
        
        doc = Document()
        
        # Add title and center it
        title = doc.add_heading('Romanian Apartment Market Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        for city in self.city_list: 
            # --- Add content to Word doc ---
            doc.add_heading(city, level=1)
            
            # Add the plot
            plot_path = os.path.join(self.images_path, f'{city}_price_per_sqm.png')
            if os.path.exists(plot_path):
                doc.add_picture(plot_path, width=Inches(6.0))
            else:
                doc.add_paragraph(f"[Plot for {city} not found at {plot_path}]")
            
            # Add the pre-calculated text
            doc.add_paragraph(self.analysis_data[city])

        # --- Add the final ranking to the Word doc ---
        doc.add_heading('Final Investment Ranking', level=1)
        
        for rank, (city, lower, upper) in enumerate(self.city_ranking, start=1):
            rank_text = f"""{rank}. {city}: 
The expected monthly income from rent for every euro spent is estimated to be between {lower:.2f} and {upper:.2f}.\n"""
            doc.add_paragraph(rank_text)

        # --- Save the document to the reports folder ---
        doc.save(full_report_path)
        print(f"Successfully saved Word report to: {full_report_path}")